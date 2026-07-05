"""GUI Streamlit de SismoAI — versión YOLOv8-seg (enfoque del profesor).

Misma interfaz que la app U-Net, pero la detección la hace un modelo YOLOv8-seg
pre-entrenado y afinado sobre Thebe. Ejecutar:
    .venv/bin/streamlit run sismoai/app_yolo.py
"""
from __future__ import annotations

import base64
import io
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

import cv2
import numpy as np
import streamlit as st
from PIL import Image, ImageDraw
from streamlit_image_coordinates import streamlit_image_coordinates

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

st.set_page_config(page_title="Proyecto Sísmica Fallas — YOLO", layout="wide")

MODEL = Path(__file__).resolve().parent.parent / "models" / "yolo_fallas.pt"
EX_DIR = Path(__file__).resolve().parent / "examples"


@st.cache_resource
def get_model():
    from ultralytics import YOLO
    return YOLO(str(MODEL))


def read_image(uploaded) -> np.ndarray:
    data = np.frombuffer(uploaded.getvalue(), np.uint8)
    return cv2.imdecode(data, cv2.IMREAD_GRAYSCALE)


def predict_instances(model, gray: np.ndarray, conf: float) -> list[dict]:
    """Devuelve una lista con cada falla detectada por separado (no fusionadas),
    ordenadas de izquierda a derecha, cada una con su máscara y su bbox."""
    rgb = cv2.cvtColor(gray, cv2.COLOR_GRAY2RGB)
    r = model.predict(rgb, imgsz=640, conf=conf, verbose=False)[0]
    H, W = gray.shape
    instancias = []
    if r.masks is not None:
        for i, m in enumerate(r.masks.data.cpu().numpy()):
            mask = cv2.resize(m, (W, H), interpolation=cv2.INTER_NEAREST) > 0.5
            ys, xs = np.where(mask)
            if len(xs) == 0:
                continue
            instancias.append({
                "id": i,
                "mask": mask,
                "bbox": (int(xs.min()), int(ys.min()), int(xs.max()), int(ys.max())),
                "cx": float(xs.mean()),
            })
    instancias.sort(key=lambda d: d["cx"])
    return instancias


def combinar_mascara(instancias: list[dict], H: int, W: int) -> np.ndarray:
    mask = np.zeros((H, W), bool)
    for inst in instancias:
        mask |= inst["mask"]
    return mask


def etiquetar_numeros(rgb_img: np.ndarray, instancias: list[dict]) -> np.ndarray:
    """Dibuja el número de cada falla (fondo amarillo) cerca de su extremo superior."""
    img = rgb_img.copy()
    for inst in instancias:
        x0, y0, _, _ = inst["bbox"]
        pos = (max(x0 - 6, 2), max(y0 - 6, 14))
        texto = str(inst["numero"])
        cv2.putText(img, texto, pos, cv2.FONT_HERSHEY_SIMPLEX, 0.6, (0, 0, 0), 3, cv2.LINE_AA)
        cv2.putText(img, texto, pos, cv2.FONT_HERSHEY_SIMPLEX, 0.6, (255, 255, 0), 1, cv2.LINE_AA)
    return img


def crear_instancia_manual(p1: tuple[int, int], p2: tuple[int, int], H: int, W: int) -> dict:
    """Crea una falla 'manual' a partir de los dos puntos que el usuario marcó con clic,
    con la misma forma de diccionario que las fallas detectadas por YOLO."""
    m = np.zeros((H, W), np.uint8)
    cv2.line(m, p1, p2, 1, thickness=4)
    mask = m.astype(bool)
    xs, ys = (p1[0], p2[0]), (p1[1], p2[1])
    return {
        "id": f"manual-{p1}-{p2}",
        "mask": mask,
        "bbox": (min(xs), min(ys), max(xs), max(ys)),
        "cx": (p1[0] + p2[0]) / 2,
        "manual": True,
    }


def overlay(gray, mask, color):
    g = gray.astype(np.float32)
    g = (g - g.min()) / (float(np.ptp(g)) or 1.0) * 255
    rgb = np.stack([g] * 3, -1).astype(np.uint8)
    rgb[cv2.dilate(mask.astype(np.uint8), np.ones((2, 2), np.uint8)) > 0] = color
    return rgb


# ---------------------------------------------------------------------------
# Reporte automático con IA (OpenAI / Gemini / DeepSeek) — Word y PDF
# ---------------------------------------------------------------------------

PROVEEDORES = {
    "OpenAI (GPT)": {"env": "OPENAI_API_KEY", "modelo": "gpt-5.4-mini"},
    "Google Gemini": {"env": "GEMINI_API_KEY", "modelo": "gemini-3.5-flash"},
    "DeepSeek": {"env": "DEEPSEEK_API_KEY", "modelo": "deepseek-v4-flash"},
}

# Solo OpenAI y Gemini tienen soporte de visión confiable por API hoy en día.
PROVEEDORES_VISION = {k: v for k, v in PROVEEDORES.items() if k != "DeepSeek"}


def get_api_key(proveedor: str) -> str | None:
    """Busca la API key primero en st.secrets (secrets.toml) y, si no está,
    en las variables de entorno del sistema."""
    nombre = PROVEEDORES[proveedor]["env"]
    try:
        if nombre in st.secrets:
            return st.secrets[nombre]
    except Exception:
        pass
    return os.environ.get(nombre)


def generar_reporte_ia(proveedor: str, prompt: str) -> str:
    """Manda el prompt al proveedor elegido y devuelve el texto del reporte."""
    api_key = get_api_key(proveedor)
    if not api_key:
        raise RuntimeError(
            f"No se encontró la API key de {proveedor} "
            f"(variable {PROVEEDORES[proveedor]['env']}). "
            "Configúrala en .streamlit/secrets.toml o como variable de entorno."
        )
    modelo = PROVEEDORES[proveedor]["modelo"]

    if proveedor == "OpenAI (GPT)":
        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        resp = client.chat.completions.create(
            model=modelo,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
        )
        return resp.choices[0].message.content

    if proveedor == "Google Gemini":
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(modelo)
        resp = model.generate_content(prompt)
        return resp.text

    if proveedor == "DeepSeek":
        from openai import OpenAI  # DeepSeek usa el SDK de OpenAI con otra base_url
        client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
        resp = client.chat.completions.create(
            model=modelo,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
        )
        return resp.choices[0].message.content

    raise ValueError(f"Proveedor no soportado: {proveedor}")


# ---------------------------------------------------------------------------
# Interpretación geológica con IA con visión (clasifica fallas + sismofacies)
# ---------------------------------------------------------------------------

def imagen_a_png_bytes(img_rgb: np.ndarray) -> bytes:
    ok, buf = cv2.imencode(".png", cv2.cvtColor(img_rgb, cv2.COLOR_RGB2BGR))
    return buf.tobytes()


def construir_prompt_vision(n_faults: int) -> str:
    return f"""Eres un geólogo estructural y estratigráfico con más de 25 años de experiencia en interpretación de sísmica de reflexión para exploración de hidrocarburos.
Tu función es analizar imágenes de secciones sísmicas 2D y generar una interpretación geológica profesional con calidad de publicación científica.
Realiza siempre una reinterpretación independiente de la imagen. No copies las anotaciones existentes; utiliza únicamente la evidencia sísmica observable para interpretar las estructuras.
Conserva exactamente la geometría, resolución, escala y dimensiones de la imagen original. No deformes, recortes ni alteres la sección sísmica.
Reinterpreta todas las fallas observadas ({n_faults} fallas detectadas automáticamente). Puedes eliminar fallas incorrectas, extender fallas incompletas, agregar nuevas fallas cuando exista evidencia sísmica suficiente, interpretar ramificaciones, fallas sintéticas, antitéticas, lístricas, de crecimiento, ciegas, normales o inversas según la geometría de los reflectores.
Las fallas principales deben representarse con líneas continuas y las fallas secundarias con líneas discontinuas.
Identifica e interpreta únicamente cuando exista evidencia sísmica clara:
• Rollovers
• Bloques rotados
• Horst
• Graben
• Half-graben
• Discordancias
• Onlap
• Downlap
• Toplap
• Truncamientos
• Canales
• Pliegues
• Colapso gravitacional
• Crestas estructurales
Genera un segundo panel con pseudo-facies geológicas.
Las pseudo-facies deben definirse exclusivamente a partir de cambios en:
- continuidad de reflectores
- amplitud
- frecuencia
- configuración interna
- geometría
- espesor
- terminaciones estratigráficas
No inventes unidades geológicas. Cada pseudo-facies debe seguir fielmente la estratigrafía observada y nunca cruzar reflectores importantes.
Utiliza colores suaves y profesionales para las pseudo-facies, manteniendo visible la información sísmica.
Incluye una leyenda discreta con:
- Falla principal
- Falla secundaria
- Cresta de rollover
- Pseudo-facies
Nunca agregues información que no exista en la imagen.
No agregues: títulos, ejes, coordenadas, cuadrículas, Norte, Sur, Este, Oeste, escalas nuevas, textos decorativos ni etiquetas innecesarias.
Conserva únicamente la barra de escala si ya está presente en la imagen original.
La interpretación debe tener la apariencia de haber sido realizada por un intérprete sísmico profesional para una publicación en revistas como AAPG Bulletin, Interpretation, Basin Research o Marine and Petroleum Geology.
Prioriza siempre la coherencia geológica sobre la cantidad de elementos interpretados. Si una estructura no es claramente identificable, no la dibujes. La interpretación debe ser técnicamente defendible y consistente con los principios de la geología estructural y la estratigrafía sísmica.

Responde ÚNICAMENTE con un JSON válido (sin texto adicional, sin bloques de código markdown), con esta estructura exacta:
{{
  "fallas": [{{"numero": 1, "tipo": "principal", "sentido": "hundido_derecha"}}],
  "sismofacies": [{{"nombre": "SF1", "y_inicio_pct": 0, "y_fin_pct": 20, "descripcion": "..."}}],
  "interpretacion": ["...", "..."]
}}"""


def interpretar_con_ia(proveedor: str, imagen_rgb: np.ndarray, n_faults: int) -> dict:
    """Manda la imagen al modelo de visión y devuelve la interpretación como diccionario."""
    api_key = get_api_key(proveedor)
    if not api_key:
        raise RuntimeError(
            f"No se encontró la API key de {proveedor} "
            f"(variable {PROVEEDORES[proveedor]['env']})."
        )
    prompt = construir_prompt_vision(n_faults)
    png_bytes = imagen_a_png_bytes(imagen_rgb)
    modelo = PROVEEDORES[proveedor]["modelo"]

    if proveedor == "OpenAI (GPT)":
        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        img_b64 = base64.b64encode(png_bytes).decode()
        resp = client.chat.completions.create(
            model=modelo,
            messages=[{"role": "user", "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_b64}"}},
            ]}],
            temperature=0.3,
        )
        texto = resp.choices[0].message.content

    elif proveedor == "Google Gemini":
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-1.5-flash")
        resp = model.generate_content([prompt, {"mime_type": "image/png", "data": png_bytes}])
        texto = resp.text

    else:
        raise ValueError(f"{proveedor} no soporta interpretación visual todavía.")

    limpio = re.sub(r"^```(json)?|```$", "", texto.strip(), flags=re.MULTILINE).strip()
    return json.loads(limpio)


def dibujar_interpretacion(gray: np.ndarray, instancias_activas: list[dict], datos: dict) -> np.ndarray:
    """Dibuja la interpretación en dos paneles: superior estructural + inferior sismofacies."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.patches as mpatches
    import matplotlib.lines as mlines
    import matplotlib.pyplot as plt
    from matplotlib.colors import to_rgba

    H, W = gray.shape
    DPI = 100
    fig_w = (W + 220) / DPI  # espacio extra para leyenda
    fig_h = (H * 2 + 20) / DPI  # dos paneles

    fig, (ax_top, ax_bot) = plt.subplots(2, 1, figsize=(fig_w, fig_h), dpi=DPI)
    fig.patch.set_facecolor("#1a1a1a")

    # ── Panel superior: interpretación estructural ──────────────────────────
    ax_top.imshow(gray, cmap="gray", extent=(0, W, H, 0), aspect="auto")
    ax_top.set_xlim(0, W)
    ax_top.set_ylim(H, 0)
    ax_top.axis("off")

    fallas_info = {f["numero"]: f for f in datos.get("fallas", [])}
    estilos_tipo = {
        "principal":             dict(lw=2.0, ls="-",  color="red"),
        "secundaria_sintetica":  dict(lw=1.2, ls="--", color="red"),
        "secundaria_antitetica": dict(lw=1.2, ls=":",  color="red"),
    }

    for inst in instancias_activas:
        numero = inst["numero"]
        info = fallas_info.get(numero, {})
        tipo = info.get("tipo", "secundaria_sintetica")
        estilo = estilos_tipo.get(tipo, estilos_tipo["secundaria_sintetica"])
        ys, xs = np.where(inst["mask"])
        if len(ys) < 2:
            continue
        coef = np.polyfit(ys, xs, 1)
        y_top_f, y_bot_f = int(ys.min()), int(ys.max())
        x_top_f, x_bot_f = np.polyval(coef, [y_top_f, y_bot_f])
        ax_top.plot([x_top_f, x_bot_f], [y_top_f, y_bot_f], **estilo)
        etiqueta = f"F{numero}" if tipo == "principal" else f"f{numero}"
        ax_top.text(x_top_f, max(y_top_f - 6, 4), etiqueta,
                    color="black", fontsize=7, fontweight="bold",
                    bbox=dict(boxstyle="round,pad=0.15", fc="#f5e642", ec="none", alpha=0.9))
        # flecha sentido hundido
        sentido = info.get("sentido", "")
        if sentido:
            dx = 12 if sentido == "hundido_derecha" else -12
            ax_top.annotate("", xy=(x_top_f + dx, y_top_f + 22), xytext=(x_top_f, y_top_f + 6),
                            arrowprops=dict(arrowstyle="->", color="darkred", lw=1.0))

    # líneas de sismofacies en panel superior
    colores_sf = ["#e8c547", "#e87b8c", "#5bc8d4", "#7ec87e", "#b87ec8", "#e89c6e"]
    for i, sf in enumerate(datos.get("sismofacies", [])):
        y0 = sf.get("y_inicio_pct", 0) / 100 * H
        color = colores_sf[i % len(colores_sf)]
        ax_top.axhline(y0, color=color, lw=1.0, ls="--", alpha=0.8)
        ax_top.text(6, max(y0 - 5, 8), sf.get("nombre", f"SF{i+1}"),
                    color=color, fontsize=7, fontweight="bold")

    # caja interpretación estructural
    puntos = datos.get("interpretacion", [])
    if puntos:
        texto_interp = "INTERPRETACIÓN ESTRUCTURAL\n" + "\n".join(
            f"• {p}" for p in puntos
        )
        ax_top.text(0.01, 0.01, texto_interp, transform=ax_top.transAxes,
                    fontsize=7.5, color="white", va="bottom", ha="left",
                    bbox=dict(boxstyle="round,pad=0.5", fc="#111111", ec="#aaaaaa", lw=0.8, alpha=0.88),
                    linespacing=1.6, fontfamily="monospace")

    # leyenda panel superior
    leyenda_handles = [
        mlines.Line2D([], [], color="red", lw=2.0, ls="-",  label="Falla principal"),
        mlines.Line2D([], [], color="red", lw=1.2, ls="--", label="Falla secundaria"),
        mlines.Line2D([], [], color="#e8c547", lw=1.0, ls="--", label="Cresta de rollover"),
    ]
    ax_top.legend(handles=leyenda_handles, loc="upper right", fontsize=6,
                  framealpha=0.75, facecolor="#111", labelcolor="white",
                  title="Elementos interpretativos", title_fontsize=6)

    # ── Panel inferior: pseudo-facies coloreadas ─────────────────────────────
    ax_bot.imshow(gray, cmap="gray", extent=(0, W, H, 0), aspect="auto", alpha=0.55)
    ax_bot.set_xlim(0, W)
    ax_bot.set_ylim(H, 0)
    ax_bot.axis("off")

    sismofacies = datos.get("sismofacies", [])
    sf_handles = []
    for i, sf in enumerate(sismofacies):
        y0 = sf.get("y_inicio_pct", 0) / 100 * H
        y1 = sf.get("y_fin_pct", 100) / 100 * H if "y_fin_pct" in sf else (
            sismofacies[i+1].get("y_inicio_pct", 100) / 100 * H if i + 1 < len(sismofacies) else H
        )
        color = colores_sf[i % len(colores_sf)]
        ax_bot.axhspan(y0, y1, color=color, alpha=0.38)
        ax_bot.axhline(y0, color=color, lw=0.8, ls="-", alpha=0.6)
        ax_bot.text(6, max(y0 + 10, 12), sf.get("nombre", f"SF{i+1}"),
                    color=color, fontsize=7, fontweight="bold",
                    bbox=dict(boxstyle="round,pad=0.1", fc="black", ec="none", alpha=0.5))
        sf_handles.append(mpatches.Patch(color=color, alpha=0.7,
                                          label=f"{sf.get('nombre', f'SF{i+1}')}: {sf.get('descripcion', '')}"))

    # fallas en panel inferior (líneas negras finas)
    for inst in instancias_activas:
        numero = inst["numero"]
        info = fallas_info.get(numero, {})
        tipo = info.get("tipo", "secundaria_sintetica")
        ys, xs = np.where(inst["mask"])
        if len(ys) < 2:
            continue
        coef = np.polyfit(ys, xs, 1)
        y_top_f, y_bot_f = int(ys.min()), int(ys.max())
        x_top_f, x_bot_f = np.polyval(coef, [y_top_f, y_bot_f])
        ls = "-" if tipo == "principal" else "--"
        lw = 1.4 if tipo == "principal" else 0.9
        ax_bot.plot([x_top_f, x_bot_f], [y_top_f, y_bot_f], color="black", lw=lw, ls=ls)

    if sf_handles:
        ax_bot.legend(handles=sf_handles, loc="upper right", fontsize=6,
                      framealpha=0.75, facecolor="#111", labelcolor="white",
                      title="Pseudo-facies", title_fontsize=6)

    fig.subplots_adjust(left=0, right=1, top=1, bottom=0, hspace=0.02)
    fig.canvas.draw()
    buf = np.asarray(fig.canvas.buffer_rgba())[:, :, :3].copy()
    plt.close(fig)
    return buf


def construir_prompt_default(image_name: str, n_faults: int, conf: float,
                              metrics: dict | None, interpretacion: dict | None = None,
                              info_proyecto: dict | None = None) -> str:
    bloque_metricas = ""
    if metrics:
        bloque_metricas = (
            f"\n- Comparación con la anotación experta: "
            f"F1={metrics['f1']:.2f}, Precisión={metrics['prec']:.2f}, Recall={metrics['rec']:.2f}"
        )

    bloque_interpretacion = ""
    if interpretacion:
        fallas_txt = "; ".join(
            f"Falla {f['numero']}: {f['tipo']} ({f['sentido']})" for f in interpretacion.get("fallas", [])
        )
        sf_txt = "; ".join(
            f"{sf['nombre']}: {sf['descripcion']}" for sf in interpretacion.get("sismofacies", [])
        )
        interp_txt = "; ".join(interpretacion.get("interpretacion", []))
        bloque_interpretacion = f"""

Interpretación geológica previa (generada por IA a partir de la imagen):
- Clasificación de fallas: {fallas_txt}
- Sismofacies identificadas: {sf_txt}
- Puntos de interpretación estructural: {interp_txt}

Usa esta interpretación como base del informe; dale coherencia narrativa, no la repitas como lista."""

    bloque_proyecto = ""
    if info_proyecto:
        campos = {
            "Proyecto": info_proyecto.get("proyecto", ""),
            "Objetivo": info_proyecto.get("objetivo", ""),
            "Cuenca": info_proyecto.get("cuenca", ""),
            "País": info_proyecto.get("pais", ""),
            "Bloque": info_proyecto.get("bloque", ""),
            "Línea sísmica": info_proyecto.get("linea_sismica", ""),
            "Tipo de dato": info_proyecto.get("tipo_dato", ""),
            "Orientación": info_proyecto.get("orientacion", ""),
            "Escala vertical": info_proyecto.get("escala_vertical", ""),
            "Intervalo interpretado": info_proyecto.get("intervalo", ""),
            "Formaciones de interés": info_proyecto.get("formaciones", ""),
            "Pozos de control": info_proyecto.get("pozos", ""),
            "Sistema tectónico": info_proyecto.get("sistema_tectonico", ""),
            "Objetivo exploratorio": info_proyecto.get("objetivo_exploratorio", ""),
            "Observaciones": info_proyecto.get("observaciones", ""),
        }
        # Solo incluir campos que el usuario llenó
        lineas = [f"- {k}: {v}" for k, v in campos.items() if v and v.strip()]
        if lineas:
            bloque_proyecto = "\n\nInformación del Proyecto (proporcionada por el usuario):\n" + "\n".join(lineas)
            bloque_proyecto += "\n\nIMPORTANTE: Usa ÚNICAMENTE la información anterior. No inventes campos ni valores que no estén listados."

    return f"""Actúa como un geocientífico senior especializado en interpretación sísmica estructural y estratigráfica para exploración de hidrocarburos, con experiencia en redacción de reportes técnicos de nivel profesional similares a los elaborados en la industria por compañías como Schlumberger, Halliburton, CGG, TGS, Petrobras o Shell.

Tu tarea es redactar un REPORTE TÉCNICO DE INTERPRETACIÓN SÍSMICA PARA EXPLORACIÓN DE HIDROCARBUROS con base en:
1) la imagen sísmica interpretada proporcionada por el usuario, y
2) la información del proyecto ingresada por el usuario, la cual puede estar completa, parcial o incompleta.

El reporte debe tener una apariencia profesional, sobria, técnica y ejecutiva, con estilo de informe geocientífico real. El objetivo es que el resultado no parezca un texto genérico de IA, sino un informe técnico elaborado por un geólogo o geofísico de exploración.

ENTRADAS DEL SISTEMA

A) Imagen sísmica interpretada:
- Imagen analizada: {image_name}
- Número de fallas detectadas: {n_faults}{bloque_metricas}{bloque_interpretacion}

B) Información del proyecto:{bloque_proyecto if bloque_proyecto else chr(10) + "No se proporcionó información del proyecto."}

ESTRUCTURA OBLIGATORIA DEL REPORTE (máximo 4 páginas):

1. INFORMACIÓN DEL PROYECTO
Presenta la información en una tabla elegante de dos columnas (Campo / Valor) y luego un párrafo técnico en prosa que resuma el proyecto.

2. RESUMEN EJECUTIVO
Texto corrido técnico y sólido. No una lista. Resume: objetivo del análisis, rasgos estructurales principales, comportamiento estratigráfico general, relevancia exploratoria, principales incertidumbres y valor del estudio.

3. FIGURA PRINCIPAL Y OBJETIVO DEL ESTUDIO
Indica dónde va la figura con este pie: "Figura 1. Interpretación sísmica-estructural generada automáticamente mediante SismoAI."
Nota: "La figura presenta la sección sísmica interpretada y los principales elementos estructurales identificados."
Luego un párrafo con el objetivo técnico del análisis.

4. CONTEXTO GEOLÓGICO
Síntesis técnica del marco geológico y petrolero del área. Si el usuario indicó cuenca o región, redacta su contexto. Si no, usa versión general honesta sin inventar detalles.

5. ANÁLISIS ESTRUCTURAL
Sección en prosa, bien desarrollada. Discute: sistema de fallas, fallas principales y secundarias, geometría de bloques, jerarquía estructural, segmentación, rollover, crecimiento estratal, compartimentación y estilo tectónico.
Puedes añadir una tabla pequeña de apoyo: Síntesis de rasgos estructurales (Elemento / Interpretación técnica / Implicación exploratoria).

6. INTERPRETACIÓN ESTRATIGRÁFICA
En prosa: continuidad de reflectores, cambios de espesor, pseudo-facies, posibles unidades reservorio, relación estratigrafía-fallamiento.
Tabla auxiliar opcional: Pseudo-facies / Interpretación preliminar / Relevancia exploratoria.

7. IMPLICACIONES PARA LA EXPLORACIÓN DE HIDROCARBUROS
Sección de alto peso técnico en texto corrido. Discute: posibles reservorios, potencial de sello, cierres contra falla, trampas estructurales, compartimentación, riesgo exploratorio, incertidumbre en continuidad lateral, relación con formaciones objetivo.
Tabla opcional: Síntesis exploratoria (Elemento del sistema / Evaluación preliminar / Incertidumbre principal).

8. CONCLUSIONES
Entre 5 y 8 conclusiones técnicas desarrolladas, no frases telegráficas. Referirse a: estilo estructural, papel de fallas, pseudo-facies, reservorio/sello/trampa, utilidad exploratoria, incertidumbres.

9. RECOMENDACIONES
Recomendaciones concretas y justificadas: integración con pozos, atributos sísmicos, inversión sísmica, análisis AVO, sísmica 3D, sellado de fallas, maduración de leads.

10. BIBLIOGRAFÍA
Referencias clásicas breves: Badley (1985), Brown (2011), Fossen (2016), Selley & Sonnenberg (2015), Sheriff & Geldart (2002), Yilmaz (2001).

REGLAS DE ESTILO:
- Tono técnico, sobrio y profesional
- Priorizar texto sobre tablas
- No mencionar que es IA ni usar frases vacías
- Redactar como informe geocientífico real
- Máximo 4 páginas, sin espacios en blanco excesivos
- Escribe en español"""


def crear_docx(texto: str, imagen: np.ndarray | None = None,
                titulo: str = "Reporte de Detección de Fallas Geológicas") -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re as _re

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    imagen_insertada = [False]  # usamos lista para modificar desde función interna

    def agregar_tabla_markdown(lineas_tabla):
        """Convierte líneas markdown de tabla a tabla Word."""
        filas = []
        for l in lineas_tabla:
            if _re.match(r"^\|[-| ]+\|$", l.strip()):
                continue
            celdas = [c.strip() for c in l.strip().strip("|").split("|")]
            filas.append(celdas)
        if not filas:
            return
        ncols = max(len(f) for f in filas)
        tabla = doc.add_table(rows=len(filas), cols=ncols)
        tabla.style = "Table Grid"
        tabla.autofit = False
        # Distribuir ancho equitativamente en el ancho disponible
        from docx.shared import Inches as _Inches
        ancho_col = int((_Inches(6.0).pt * 914400 / 914400) / ncols * 914400)
        for col in tabla.columns:
            for cell in col.cells:
                cell.width = _Inches(6.0 / ncols)
        for i, fila in enumerate(filas):
            for j, celda in enumerate(fila):
                if j < ncols:
                    cell = tabla.cell(i, j)
                    cell.text = celda
                    # Primera fila en negrita
                    if i == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
        doc.add_paragraph("")

    lineas = texto.split("\n")
    buffer_tabla = []
    i = 0
    while i < len(lineas):
        linea = lineas[i]
        stripped = linea.strip()

        # Detectar tabla markdown
        if stripped.startswith("|"):
            buffer_tabla.append(stripped)
            i += 1
            continue
        else:
            if buffer_tabla:
                agregar_tabla_markdown(buffer_tabla)
                buffer_tabla = []

        # Encabezados
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:-2])
            run.bold = True
        elif stripped.startswith("- ") or stripped.startswith("• "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped == "---" or stripped == "***":
            pass  # separador, ignorar
        elif stripped:
            doc.add_paragraph(stripped)
            # Insertar imagen justo después de la línea "Figura 1"
            if imagen is not None and not imagen_insertada[0] and "Figura 1" in stripped:
                doc.add_paragraph("")
                img_buf = io.BytesIO(imagen_a_png_bytes(imagen))
                doc.add_picture(img_buf, width=Inches(6.0))
                doc.add_paragraph("")
                imagen_insertada[0] = True
        i += 1

    # Si no se insertó la imagen (no había "Figura 1"), insertarla al final
    if imagen is not None and not imagen_insertada[0]:
        doc.add_paragraph("")
        img_buf = io.BytesIO(imagen_a_png_bytes(imagen))
        doc.add_picture(img_buf, width=Inches(6.0))

    if buffer_tabla:
        agregar_tabla_markdown(buffer_tabla)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def crear_pdf(texto: str, imagen: np.ndarray | None = None,
              titulo: str = "Reporte de Detección de Fallas Geológicas") -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import Image as RLImage
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    import re as _re

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER,
                            leftMargin=1.2*inch, rightMargin=1.2*inch,
                            topMargin=1*inch, bottomMargin=1*inch)
    styles = getSampleStyleSheet()

    # Estilos personalizados
    estilo_titulo = ParagraphStyle("Titulo", parent=styles["Title"], fontSize=16, spaceAfter=12, alignment=1)
    estilo_h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=14, spaceAfter=6, spaceBefore=12)
    estilo_h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=12, spaceAfter=4, spaceBefore=10)
    estilo_h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=11, spaceAfter=4, spaceBefore=8)
    estilo_normal = ParagraphStyle("Normal2", parent=styles["Normal"], fontSize=10, spaceAfter=6, leading=14)
    estilo_bullet = ParagraphStyle("Bullet", parent=styles["Normal"], fontSize=10, leftIndent=20, spaceAfter=4, bulletIndent=10)
    estilo_meta = ParagraphStyle("Meta", parent=styles["Normal"], fontSize=9, textColor=colors.grey, spaceAfter=12)

    elementos = []
    imagen_insertada = [False]

    def procesar_tabla_markdown(lineas_tabla):
        filas = []
        for l in lineas_tabla:
            if _re.match(r"^\|[-| ]+\|$", l.strip()):
                continue
            celdas = [c.strip() for c in l.strip().strip("|").split("|")]
            filas.append(celdas)
        if not filas:
            return None
        ncols = max(len(f) for f in filas)
        # Normalizar filas
        filas = [f + [""] * (ncols - len(f)) for f in filas]
        # Ancho disponible: LETTER - márgenes
        ancho_disponible = 6.1 * inch
        col_widths = [ancho_disponible / ncols] * ncols

        # Convertir celdas a Paragraph para que el texto haga wrap automático
        estilo_celda = ParagraphStyle("celda", fontSize=9, leading=12, wordWrap="CJK")
        estilo_celda_header = ParagraphStyle("celda_h", fontSize=9, leading=12,
                                              textColor=colors.white, fontName="Helvetica-Bold", wordWrap="CJK")
        filas_p = []
        for i_fila, fila in enumerate(filas):
            fila_p = []
            for celda in fila:
                estilo = estilo_celda_header if i_fila == 0 else estilo_celda
                fila_p.append(Paragraph(celda, estilo))
            filas_p.append(fila_p)

        tabla = Table(filas_p, hAlign="LEFT", colWidths=col_widths, repeatRows=1)
        tabla.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f2f2f2")]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        return tabla

    lineas = texto.split("\n")
    buffer_tabla = []
    i = 0
    while i < len(lineas):
        linea = lineas[i]
        stripped = linea.strip()

        if stripped.startswith("|"):
            buffer_tabla.append(stripped)
            i += 1
            continue
        else:
            if buffer_tabla:
                t = procesar_tabla_markdown(buffer_tabla)
                if t:
                    elementos.append(t)
                    elementos.append(Spacer(1, 8))
                buffer_tabla = []

        if stripped.startswith("### "):
            elementos.append(Paragraph(stripped[4:].strip(), estilo_h3))
        elif stripped.startswith("## "):
            elementos.append(Paragraph(stripped[3:].strip(), estilo_h2))
        elif stripped.startswith("# "):
            elementos.append(Paragraph(stripped[2:].strip(), estilo_h1))
        elif stripped.startswith("- ") or stripped.startswith("• "):
            elementos.append(Paragraph(f"• {stripped[2:]}", estilo_bullet))
        elif stripped == "---":
            elementos.append(Spacer(1, 8))
        elif stripped:
            elementos.append(Paragraph(stripped, estilo_normal))
            # Insertar imagen justo después de "Figura 1"
            if imagen is not None and not imagen_insertada[0] and "Figura 1" in stripped:
                elementos.append(Spacer(1, 8))
                ancho = 6.0 * inch
                alto = ancho * imagen.shape[0] / imagen.shape[1]
                elementos.append(RLImage(io.BytesIO(imagen_a_png_bytes(imagen)), width=ancho, height=alto))
                elementos.append(Spacer(1, 12))
                imagen_insertada[0] = True
        else:
            elementos.append(Spacer(1, 4))
        i += 1

    if buffer_tabla:
        t = procesar_tabla_markdown(buffer_tabla)
        if t:
            elementos.append(t)

    # Si no se insertó la imagen, agregarla al final
    if imagen is not None and not imagen_insertada[0]:
        ancho = 6.0 * inch
        alto = ancho * imagen.shape[0] / imagen.shape[1]
        elementos.append(RLImage(io.BytesIO(imagen_a_png_bytes(imagen)), width=ancho, height=alto))

    doc.build(elementos)
    return buf.getvalue()


st.title("Proyecto Sísmica Fallas — YOLO")

with st.sidebar:
    st.header("Parámetros")
    conf = st.slider("Confianza mínima", 0.05, 0.9, 0.10, 0.05,
                     help="Qué tan seguro debe estar YOLO para marcar una falla. "
                          "Más bajo = detecta más fallas (mejor para líneas finas).")
    try:
        model = get_model()
        st.success("Modelo YOLOv8-seg cargado")
        st.caption("Datos: Thebe (real)")
    except Exception as e:  # noqa
        st.error(f"No se pudo cargar el modelo YOLO: {e}")
        st.stop()

st.caption("Detección de fallas con **YOLOv8-seg** (pre-entrenado + fine-tuning sobre Thebe). "
           "Sube una sección sísmica en gris o prueba un ejemplo real.")

ejemplos = sorted(p for p in EX_DIR.glob("thebe_*.png") if "_gt" not in p.name)
nombres = ["— ninguno —"] + [f"Ejemplo {i + 1} (Thebe real)" for i in range(len(ejemplos))]

uploaded = st.file_uploader("Sube tu imagen sísmica (PNG / JPG / TIFF)",
                            type=["png", "jpg", "jpeg", "tif", "tiff"])
sel = st.selectbox("…o prueba con una sección REAL del dataset Thebe "
                   "(incluye la anotación de los geólogos para comparar):", nombres)

gt_mask = None
if uploaded is not None:
    gray = read_image(uploaded)
    image_name = uploaded.name
elif sel != "— ninguno —":
    p = ejemplos[nombres.index(sel) - 1]
    gray = cv2.imread(str(p), cv2.IMREAD_GRAYSCALE)
    image_name = p.name
    gtp = p.with_name(p.stem + "_gt.png")
    if gtp.exists():
        gt_mask = cv2.imread(str(gtp), cv2.IMREAD_GRAYSCALE) > 127
    st.caption("Sección real de Thebe — con anotación experta para comparar.")
else:
    st.info("⬆️ Sube una imagen, **o** elige un ejemplo real del dataset Thebe.")
    st.stop()

if gray is None:
    st.error("No se pudo leer la imagen.")
    st.stop()

H, W = gray.shape
with st.spinner("Detectando fallas con YOLO..."):
    instancias = predict_instances(model, gray, conf)
for idx, inst in enumerate(instancias, start=1):
    inst["numero"] = idx

if not instancias:
    st.warning("No se detectaron fallas con el umbral de confianza actual. Prueba bajando el slider.")
    st.image(gray, clamp=True, use_container_width=True)
    st.stop()

mask_preview = combinar_mascara(instancias, H, W)
preview_img = etiquetar_numeros(overlay(gray, mask_preview, (255, 0, 0)), instancias)
st.image(preview_img, caption=f"Detección preliminar — {len(instancias)} fallas numeradas",
         use_container_width=True)

opciones_fallas = [f"Falla {inst['numero']}" for inst in instancias]
a_eliminar = st.multiselect(
    "¿Alguna marca es un falso positivo? Desmárcala aquí antes de continuar:",
    opciones_fallas, key="fallas_a_eliminar",
)
ids_excluidos = {inst["id"] for inst in instancias if f"Falla {inst['numero']}" in a_eliminar}
instancias_filtradas = [inst for inst in instancias if inst["id"] not in ids_excluidos]

# ---------------------------------------------------------------------------
# Agregar fallas manuales por clic
# ---------------------------------------------------------------------------
st.markdown("##### ✏️ Agregar una falla manual")
st.caption("¿Ves una falla que YOLO no marcó? Haz clic en su extremo superior y luego en su extremo "
           "inferior, sobre la imagen de abajo, y confírmala con el botón.")

if "fallas_manuales" not in st.session_state:
    st.session_state["fallas_manuales"] = []
if "puntos_temp" not in st.session_state:
    st.session_state["puntos_temp"] = []
if "ultimo_click" not in st.session_state:
    st.session_state["ultimo_click"] = None

mask_intermedia = combinar_mascara(instancias_filtradas, H, W)
overlay_intermedio = etiquetar_numeros(overlay(gray, mask_intermedia, (255, 0, 0)), instancias_filtradas)

click = streamlit_image_coordinates(Image.fromarray(overlay_intermedio), key="click_falla_manual")

if click is not None:
    firma = click.get("time", (click.get("x"), click.get("y")))
    if firma != st.session_state["ultimo_click"]:
        st.session_state["ultimo_click"] = firma
        st.session_state["puntos_temp"].append((int(click["x"]), int(click["y"])))
        st.session_state["puntos_temp"] = st.session_state["puntos_temp"][-2:]

puntos_temp = st.session_state["puntos_temp"]
if len(puntos_temp) == 1:
    st.info(f"Primer punto marcado en {puntos_temp[0]}. Haz clic en el segundo punto (extremo inferior).")
elif len(puntos_temp) == 2:
    p1, p2 = puntos_temp
    cc1, cc2 = st.columns(2)
    if cc1.button(f"➕ Agregar falla {p1} → {p2}"):
        st.session_state["fallas_manuales"].append({"p1": p1, "p2": p2})
        st.session_state["puntos_temp"] = []
        st.rerun()
    if cc2.button("🔄 Reiniciar selección"):
        st.session_state["puntos_temp"] = []
        st.rerun()

if st.session_state["fallas_manuales"]:
    st.caption("Fallas manuales agregadas:")
    for i, m in enumerate(st.session_state["fallas_manuales"]):
        fcol1, fcol2 = st.columns([5, 1])
        fcol1.write(f"Manual {i + 1}: {m['p1']} → {m['p2']}")
        if fcol2.button("🗑️", key=f"borrar_manual_{i}"):
            st.session_state["fallas_manuales"].pop(i)
            st.rerun()

manuales_instancias = [crear_instancia_manual(m["p1"], m["p2"], H, W)
                        for m in st.session_state["fallas_manuales"]]
instancias_activas = instancias_filtradas + manuales_instancias
for idx, inst in enumerate(instancias_activas, start=1):
    inst["numero"] = idx

n_faults = len(instancias_activas)
mask = combinar_mascara(instancias_activas, H, W)
pred_ov = etiquetar_numeros(overlay(gray, mask, (255, 0, 0)), instancias_activas)

if gt_mask is not None and gt_mask.shape == mask.shape:
    real_ov = overlay(gray, gt_mask, (0, 200, 0))
    c1, c2, c3 = st.columns(3)
    c1.subheader("Original"); c1.image(gray, clamp=True, use_container_width=True)
    c2.subheader("🟢 Real (geólogos)"); c2.image(real_ov, use_container_width=True)
    c3.subheader("🔴 Predicción YOLO"); c3.image(pred_ov, use_container_width=True)
    k = np.ones((9, 9), np.uint8)
    gt_d = cv2.dilate(gt_mask.astype(np.uint8), k) > 0
    pr_d = cv2.dilate(mask.astype(np.uint8), k) > 0
    prec = int((mask & gt_d).sum()) / int(mask.sum()) if mask.sum() else 0.0
    rec = int((gt_mask & pr_d).sum()) / int(gt_mask.sum()) if gt_mask.sum() else 0.0
    f1 = 2 * prec * rec / (prec + rec) if (prec + rec) else 0.0
    st.markdown("##### Comparación con la anotación real (con tolerancia)")
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Fallas detectadas", n_faults)
    m2.metric("Coincidencia (F1)", f"{f1:.2f}")
    m3.metric("Precisión", f"{prec:.2f}")
    m4.metric("Recall", f"{rec:.2f}")
    st.caption("🟢 verde = geólogos · 🔴 rojo = YOLO. Métricas con tolerancia (±9 px).")
    metrics = {"f1": f1, "prec": prec, "rec": rec}
else:
    col1, col2 = st.columns(2)
    col1.subheader("Original"); col1.image(gray, clamp=True, use_container_width=True)
    col2.subheader("Detección YOLO (fallas en rojo)"); col2.image(pred_ov, use_container_width=True)
    st.metric("Fallas geológicas detectadas", n_faults)
    metrics = None

# ---------------------------------------------------------------------------
img_con_sf = pred_ov



# ---------------------------------------------------------------------------
# Apartado: Interpretación geológica con IA (visión)
# ---------------------------------------------------------------------------
st.divider()
st.header("🔎 Interpretación geológica con IA")

# --- Advertencia y opción alternativa ---
st.warning(
    "⚠️ **Advertencia:** La interpretación generada por IA en esta sección tiene limitaciones técnicas. "
    "Las sismofacies se representan como bandas horizontales aproximadas y no siguen los reflectores sísmicos reales. "
    "Se recomienda usarla únicamente como referencia preliminar y bajo criterio del intérprete."
)

with st.expander("💡 ¿Quieres una interpretación más precisa? Hazlo tú mismo con ChatGPT o Gemini", expanded=False):
    st.markdown("""
### Pasos para generar tu propia interpretación de alta calidad

**1. Descarga la imagen de detección YOLO** (la que aparece arriba con las fallas en rojo).

**2. Ve a una de estas plataformas:**
- 🤖 **ChatGPT:** [https://chat.openai.com](https://chat.openai.com)
- 🌐 **Gemini:** [https://gemini.google.com](https://gemini.google.com)

**3. Sube la imagen y usa el siguiente prompt:**

```
Eres un experto intérprete en geología estructural y sísmica de reflexión.
Se te pide que realices un análisis de la siguiente imagen, generando una imagen estética
de las fallas coloreadas en rojo sin inventarte nuevas. También debes interpretar qué tipo
de fallas son. No debes utilizar datos que no estén en la imagen.
También se te pide que dibujes pseudo-facies siguiendo los reflectores, generando una imagen
interpretada. Estas pseudo-facies serán identificadas con números (SF1, SF2, SF3...), 
no poner litologías ni nombres geológicos inventados.
Debes generar una interpretación del posible ambiente geológico y las posibles estructuras
presentes en la imagen. No pongas direcciones geográficas como E-W, N-S, etc.
La imagen final debe verse profesional, con leyenda, y similar a una publicación científica.
```

> 📝 Puedes modificar el prompt según tu línea sísmica y contexto geológico.

**4. Descarga la imagen generada por la IA.**

**5. Sube esa imagen aquí abajo** para usarla en el reporte automático:
""")
    imagen_externa = st.file_uploader(
        "📁 Sube aquí tu imagen interpretada (de ChatGPT o Gemini)",
        type=["png", "jpg", "jpeg"],
        key="imagen_externa"
    )
    if imagen_externa is not None:
        img_ext = np.array(Image.open(imagen_externa).convert("RGB"))
        st.session_state["imagen_interpretada"] = img_ext
        st.image(img_ext, caption="Imagen interpretada cargada correctamente ✅", use_container_width=True)
        st.success("¡Imagen cargada! Esta se usará en el reporte automático.")

st.caption("O usa la interpretación automática integrada (con las limitaciones indicadas arriba).")

proveedor_vision = st.selectbox("Proveedor de IA (debe soportar visión)",
                                list(PROVEEDORES_VISION.keys()), key="proveedor_vision")

if st.button("🔎 Interpretar con IA"):
    with st.spinner(f"Analizando la imagen con {proveedor_vision}..."):
        try:
            datos = interpretar_con_ia(proveedor_vision, img_con_sf, n_faults)
            st.session_state["datos_interpretacion"] = datos
            st.session_state["imagen_interpretada"] = dibujar_interpretacion(gray, instancias_activas, datos)
        except Exception as e:  # noqa
            st.session_state.pop("datos_interpretacion", None)
            st.session_state.pop("imagen_interpretada", None)
            st.error(f"No se pudo interpretar la imagen: {e}")

if "imagen_interpretada" in st.session_state:
    st.image(st.session_state["imagen_interpretada"],
             caption="Interpretación geológica generada por IA", use_container_width=True)
    if st.session_state.get("datos_interpretacion"):
        with st.expander("Ver datos crudos de la interpretación (JSON)"):
            st.json(st.session_state["datos_interpretacion"])

# ---------------------------------------------------------------------------
# Apartado: Reporte automático con IA
# ---------------------------------------------------------------------------
st.divider()
st.header("📝 Reporte automático con IA")
st.caption("Genera un informe técnico redactado por IA a partir de los resultados de la detección, "
           "y descárgalo en Word o PDF.")

proveedor = st.selectbox("Proveedor de IA", list(PROVEEDORES.keys()), key="proveedor_ia")

# --- Información del Proyecto ---
with st.expander("📋 Información del Proyecto (opcional pero recomendado)", expanded=True):
    st.caption("Completa los datos de tu proyecto para que el reporte sea más preciso y contextualizado.")
    col1, col2 = st.columns(2)
    with col1:
        ip_proyecto = st.text_input("Proyecto", placeholder="Ej: Campo Jivino")
        ip_objetivo = st.text_input("Objetivo", placeholder="Ej: Evaluación estructural para exploración de hidrocarburos")
        ip_cuenca = st.text_input("Cuenca", placeholder="Ej: Cuenca Oriente")
        ip_pais = st.text_input("País", placeholder="Ej: Ecuador")
        ip_bloque = st.text_input("Bloque", placeholder="Ej: 15")
        ip_linea = st.text_input("Línea sísmica", placeholder="Ej: JV-24")
        ip_tipo_dato = st.text_input("Tipo de dato", placeholder="Ej: Sísmica de reflexión 2D")
    with col2:
        ip_orientacion = st.text_input("Orientación", placeholder="Ej: SW-NE")
        ip_escala = st.text_input("Escala vertical", placeholder="Ej: Tiempo (ms TWT)")
        ip_intervalo = st.text_input("Intervalo interpretado", placeholder="Ej: 0–3500 ms")
        ip_formaciones = st.text_input("Formaciones de interés", placeholder="Ej: Hollín, Napo y Tena")
        ip_pozos = st.text_input("Pozos de control", placeholder="Ej: Jivino-01 y Jivino-03")
        ip_sistema = st.text_input("Sistema tectónico", placeholder="Ej: Extensional")
        ip_obj_exp = st.text_input("Objetivo exploratorio", placeholder="Ej: Identificar fallas y trampas estructurales")
    ip_observaciones = st.text_area("Observaciones", placeholder="Ej: La sección atraviesa un sistema de fallas normales asociado al borde occidental de la cuenca.", height=80)

info_proyecto = {
    "proyecto": ip_proyecto,
    "objetivo": ip_objetivo,
    "cuenca": ip_cuenca,
    "pais": ip_pais,
    "bloque": ip_bloque,
    "linea_sismica": ip_linea,
    "tipo_dato": ip_tipo_dato,
    "orientacion": ip_orientacion,
    "escala_vertical": ip_escala,
    "intervalo": ip_intervalo,
    "formaciones": ip_formaciones,
    "pozos": ip_pozos,
    "sistema_tectonico": ip_sistema,
    "objetivo_exploratorio": ip_obj_exp,
    "observaciones": ip_observaciones,
} 

# Guardar info_proyecto en session_state para que no se pierda al hacer clic
st.session_state["info_proyecto"] = info_proyecto

interpretacion_actual = st.session_state.get("datos_interpretacion")

prompt_usuario = st.text_area(
    "Prompt para la IA (puedes editarlo antes de generar el reporte)",
    value=construir_prompt_default(image_name, n_faults, conf, metrics, interpretacion_actual, info_proyecto),
    height=280,
    key="prompt_reporte",
)
if interpretacion_actual:
    st.caption("✅ Se incluyó la interpretación geológica de IA generada arriba en este prompt.")

if st.button("🪄 Generar reporte automático"):
    with st.spinner(f"Generando reporte con {proveedor}..."):
        try:
            # Construir el prompt fresco con la info actual del proyecto
            prompt_final = construir_prompt_default(
                image_name, n_faults, conf, metrics,
                st.session_state.get("datos_interpretacion"),
                info_proyecto
            )
            st.session_state["reporte_texto"] = generar_reporte_ia(proveedor, prompt_final)
        except Exception as e:  # noqa
            st.session_state.pop("reporte_texto", None)
            st.error(f"No se pudo generar el reporte: {e}")

if "reporte_texto" in st.session_state:
    st.markdown("#### Reporte generado")
    st.write(st.session_state["reporte_texto"])

    imagen_para_reporte = st.session_state.get("imagen_interpretada", pred_ov)
    docx_bytes = crear_docx(st.session_state["reporte_texto"], imagen_para_reporte)
    pdf_bytes = crear_pdf(st.session_state["reporte_texto"], imagen_para_reporte)

    dcol1, dcol2 = st.columns(2)
    dcol1.download_button(
        "⬇️ Descargar en Word (.docx)", data=docx_bytes,
        file_name=f"reporte_fallas_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    dcol2.download_button(
        "⬇️ Descargar en PDF (.pdf)", data=pdf_bytes,
        file_name=f"reporte_fallas_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
        mime="application/pdf",
    )
